from django.shortcuts import render, get_object_or_404, redirect
from django.http import HttpResponse, Http404
import mimetypes
import os, django
import pandas as pd
from find.forms import EmailForm
os.environ.setdefault("DJANGO_SETTINGS_MODULE", "config.settings")# project_name project name
django.setup()
from find.models import Keyword, Result
import urllib
from django.views.decorators.csrf import csrf_exempt
import threading
from win32com.client import Dispatch
import os
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from find.models import Keyword
from find.forms import KeywordForm
from find import Search
import pythoncom
import smtplib  # 메일을 보내기 위한 라이브러리 모듈
from email.mime.multipart import MIMEMultipart
from email.mime.base import MIMEBase
from email.mime.text import MIMEText


def home(request):
    '''
        웹페이지 홈화면 제공
    '''
    return render(request, 'home.html')


def save_keyword(request):
    '''
        웹페이지에서 입력받은 키워드로
        1. 다른 함수에 전달할 키워드 가공
        2. 언론사별 키워드 검색결과 크롤링
        3. 검색결과 파일저장
        4. 웹페이지에 검색결과 내용전달
    '''
    if request.method == "POST":
        # 웹페이지에서 입력받은 키워드를 데이터베이스에 저장
        kwd = KeywordForm(request.POST)
        keyword = request.POST['keyword']
        # print(keyword)
        kwd.save()

        # 검색한 키워드의 id 값 반환
        #     : 가장 최근 저장된 데이터의 id = 저장된 데이터 개수
        kwd_id = Keyword.objects.count()

        # 언론사별 크롤링 수행
        mh = Search.munhwa(keyword)
        mh_news = mh[0]
        mh_title = mh[1]
        mh_date = mh[2]
        mh_link = mh[3]
        mh_summary = mh[4]

        hkr = Search.hankyoreh(keyword)
        hkr_news = hkr[0]
        hkr_title = hkr[1]
        hkr_date = hkr[2]
        hkr_link = hkr[3]
        hkr_summary = hkr[4]

        khn = Search.khann(keyword)
        khn_news = khn[0]
        khn_title = khn[1]
        khn_date = khn[2]
        khn_link = khn[3]
        khn_summary = khn[4]

        omy = Search.ohmyy(keyword)
        omy_news = omy[0]
        omy_title = omy[1]
        omy_date = omy[2]
        omy_link = omy[3]
        omy_summary = omy[4]

        da = Search.DongA(keyword)
        da_news = da[0]
        da_title = da[1]
        da_date = da[2]
        da_link = da[3]
        da_summary = da[4]

        ja = Search.jungang(keyword)
        ja_news = ja[0]
        ja_title = ja[1]
        ja_date = ja[2]
        ja_link = ja[3]
        ja_summary = ja[4]

        df_list = [mh, da, ja, hkr, khn, omy]
        columns = ['언론사', '제목', '날짜', 'URL', '요약']

        # 검색결과에 대한 데이터프레임 생성
        df = pd.DataFrame(df_list, columns=columns)
        # print(df)

        # 데이터베이스에 검색결과 저장
        Search.save_result(df, kwd_id)

        # Sub-Thread로 word, pdf 파일 생성
        threading.Timer(1, mk_word, (df, keyword)).start()

        # 검색결과 화면에 내용전달
        return render(request, 'result.html',
                      {'mh_news': mh_news, 'mh_title': mh_title, 'mh_date': mh_date, 'mh_link': mh_link,
                       'mh_summary': mh_summary,
                       'hkr_news': hkr_news, 'hkr_title': hkr_title, 'hkr_date': hkr_date, 'hkr_link': hkr_link,
                       'hkr_summary': hkr_summary,
                       'khn_news': khn_news, 'khn_title': khn_title, 'khn_date': khn_date, 'khn_link': khn_link,
                       'khn_summary': khn_summary,
                       'omy_news': omy_news, 'omy_title': omy_title, 'omy_date': omy_date, 'omy_link': omy_link,
                       'omy_summary': omy_summary,
                       'da_news': da_news, 'da_title': da_title, 'da_date': da_date, 'da_link': da_link,
                       'da_summary': da_summary,
                       'ja_news': ja_news, 'ja_title': ja_title, 'ja_date': ja_date, 'ja_link': ja_link,
                       'ja_summary': ja_summary, 'keyword': keyword})


def mk_word(df, keyword):
    '''
        word , pdf 파일 저장
        저장위치 : 프로젝트 폴더(최상위 폴더)
    '''
    doc = Document()
    # 문서 제목
    tit = doc.add_heading("", level=0)
    tit.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    tit.add_run('기사 종합 보고서').bold = True

    rpt_kwd = doc.add_paragraph("")
    rpt_kwd.add_run('검색어 : ' + keyword).bold = True
    rpt_kwd.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

    for i in range(len(df)):
        # 기사제목
        doc.add_heading(df.iloc[i][1], level=1)
        # 신문사
        p = doc.add_paragraph(df.iloc[i][0])
        # 날짜
        p.add_run("                 " + str(df.iloc[i][2]))
        p.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
        # 내용
        doc.add_paragraph(df.iloc[i][4])
        # link
        doc.add_paragraph(df.iloc[i][3] + '\n')
    # word 저장
    fname = 'search_result_' + keyword
    doc.save('./word/' + fname + '.docx')
    print('DOCX 저장완료')

    '''
        PDF 생성 및 저장
    '''
    pythoncom.CoInitialize()
    wordapp = Dispatch("Word.Application")
    fpath = os.path.join(os.path.abspath(".//word"), fname + '.docx')
    testDoc = wordapp.Documents.Open(FileName=fpath)

    pdfpath = os.path.join(os.path.abspath(".//word"), fname + '.pdf')
    testDoc.SaveAs(pdfpath, FileFormat=17)

    testDoc.Close()
    wordapp.Quit()
    pythoncom.CoUninitialize()
    print('PDF 저장완료')


# if __name__=='__main__':
